#!/usr/bin/env python3
"""
Escritores compartidos de tabla para los productos T01-T03 del paso 12.

Cada tabla conserva su propio script con su contrato, su fuente y sus
aserciones de entrada. Lo que se comparte aquí es únicamente la mecánica de
escritura, para que las tres salgan con el mismo aspecto: tres reglas
horizontales, ninguna vertical, unidades en el encabezado y cifras
significativas fijadas por columna.

Formatos generados por tabla: .docx (destino real, el informe se maqueta en
Word), .md (revisión rápida), .csv (dato) y .tex (booktabs, por si el trabajo
migra a LaTeX).
"""

from __future__ import annotations

import sys
from pathlib import Path

import pandas as pd

sys.path.insert(0, str(Path(__file__).resolve().parent))
import sciviz

COMA_DECIMAL = True


def _fmt(valor, decimales: int | None) -> str:
    if decimales is None:
        return str(valor)
    texto = f"{valor:,.{decimales}f}".replace(",", " ")  # espacio fino de millar
    return texto.replace(".", ",") if COMA_DECIMAL else texto


def _encabezados(df: pd.DataFrame, unidades: dict[str, str]) -> list[str]:
    return [f"{c} ({unidades[c]})" if c in unidades else c for c in df.columns]


def a_markdown(df, caption, unidades, decimales, notas) -> str:
    d = df.copy()
    for c, k in decimales.items():
        if c in d:
            d[c] = d[c].map(lambda v: _fmt(v, k))
    d.columns = _encabezados(df, unidades)
    d = d.astype(str)
    lineas = [f"**{caption}**", "", d.to_markdown(index=False), ""]
    lineas += [f"*{n}*" for n in notas]
    return "\n".join(lineas)


def a_docx(df, caption, unidades, decimales, notas, ruta: Path):
    """Salida Word con el aspecto booktabs: tres reglas horizontales, ninguna
    vertical. Una tabla pegada desde Excel se reconoce justo por lo contrario."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def borde(celda, lado, grosor):
        tcPr = celda._tc.get_or_add_tcPr()
        bordes = tcPr.find(qn("w:tcBorders"))
        if bordes is None:
            bordes = OxmlElement("w:tcBorders")
            tcPr.append(bordes)
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "single" if grosor else "nil")
        el.set(qn("w:sz"), str(grosor))
        el.set(qn("w:color"), "000000")
        bordes.append(el)

    doc = Document()
    p = doc.add_paragraph()
    p.add_run(caption).bold = True

    encabezados = _encabezados(df, unidades)
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    for j, h in enumerate(encabezados):
        celda = t.rows[0].cells[j]
        celda.text = h
        celda.paragraphs[0].runs[0].bold = True

    for _, fila in df.iterrows():
        celdas = t.add_row().cells
        for j, c in enumerate(df.columns):
            numerica = c in decimales
            celdas[j].text = _fmt(fila[c], decimales.get(c)) if numerica else str(fila[c])
            if numerica:
                celdas[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    n = len(t.rows)
    for i, fila in enumerate(t.rows):
        for celda in fila.cells:
            borde(celda, "left", 0)
            borde(celda, "right", 0)
            borde(celda, "top", 8 if i == 0 else (4 if i == 1 else 0))
            borde(celda, "bottom", 8 if i == n - 1 else (4 if i == 0 else 0))
            for par in celda.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)

    for nota in notas:
        par = doc.add_paragraph(nota)
        par.runs[0].font.size = Pt(8)

    doc.save(ruta)


def escribir(df, *, outdir: Path, stem: str, caption: str, label: str,
             unidades: dict, decimales: dict, notas: list[str]) -> None:
    outdir.mkdir(parents=True, exist_ok=True)
    (outdir / f"{stem}.tex").write_text(
        sciviz.booktabs_latex(df, caption, label, units=unidades,
                              decimals=decimales, notes=notas),
        encoding="utf-8")
    (outdir / f"{stem}.md").write_text(
        a_markdown(df, caption, unidades, decimales, notas), encoding="utf-8")
    df.to_csv(outdir / f"{stem}.csv", index=False)
    a_docx(df, caption, unidades, decimales, notas, outdir / f"{stem}.docx")

    print(f"{stem}: {len(df)} filas x {len(df.columns)} columnas "
          f"-> docx, md, csv, tex")
    if len(df) > 10 or len(df.columns) > 7:
        print("  aviso: pasada la escala ~10x7 el lector compara en vez de "
              "consultar; valorar si corresponde una figura")
